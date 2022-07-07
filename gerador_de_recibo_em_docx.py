print("Porque para mim o viver é Cristo, e o morrer é ganho. Filipenses 1:21")

import tkinter as tk

from tkinter import *
from tkinter import ttk
from tkinter import tix

import docx
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

root = tix.Tk()

class funcoes():

    # Menu de Opções
    def menus(self):
        menu_de_opcoes = Menu(self.root)
        self.root.config(menu=menu_de_opcoes)

        file_opcoes = Menu(menu_de_opcoes)

        menu_de_opcoes.add_cascade(label='Opções', menu=file_opcoes)

        file_opcoes.add_command(label='Sair', command="")

    # Opção Sair
    def sair(self):
        print("\nSaindo do programa...")
        quit()

    # Variáveis do Gerador de Arquivo

    def gerar_recibo(self):

        # 0.0 Variáveis

        i_valor_numerico = self.valor_numerico_do_recibo.get()
        i_valor_extenso = self.valor_por_extenso_do_recibo.get()

        i_pagador = self.pagador.get()
        i_cpf = self.cpf.get()
        i_apto = self.apartamento.get()
        i_valor_total_n = self.valor_numerico_total.get()
        i_valor_total_e = self.valor_por_extenso_total.get()

        i_dia = self.dia.get()
        i_mes = self.mes.get()
        i_ano = self.ano.get()

        i_impressao = self.imp.get()

        # 0.1 Configurações do Documento

        document = docx.Document()

        # 1.0 Cabeçalho

        section_h = document.sections[0]
        header = section_h.header

        # 1.1 Cabeçalho: Margens

        paragraph_h = header.paragraphs[0]
        header = section_h.header
        section_h.header_distance = Inches(0.1)
        paragraph_h.paragraph_format.left_indent = -Inches(1.18)

        # 1.1 Cabeçalho: Imagem

        run_h = paragraph_h.add_run()
        # run_h.add_picture("### ENDEREÇO DA IMAGEM ###")

        # 2.0 Título: RECIBO

        recibo = document.add_paragraph()
        recibo.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        recibo = recibo.add_run('\nRECIBO\n')
        fonte_do_recibo = recibo.font
        fonte_do_recibo.name = 'Calibri'
        fonte_do_recibo.size = Pt(36)

        # 3.0 Subtítulo: Valor - Entrada(s)

        print('\nValor Numérico:', i_valor_numerico)
        print()

        print('Valor por Extenso:', i_valor_extenso)
        print()

        # 3.1 Subtítulo: Valor

        valor = document.add_paragraph()
        valor.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        valor = valor.add_run(
            '\nVALOR -----------------------------------------------------------------------R$ ' + i_valor_numerico + '\n')
        valor.bold = True
        fonte_do_valor = valor.font
        fonte_do_valor.name = 'Calibri'
        fonte_do_valor.size = Pt(14)

        # 4.0 Texto - Entrada(s)

        print('Nome do Pagador:', i_pagador)
        print()

        print('CPF do Pagador:', i_cpf)
        print()

        print('Número do Apartamento:', i_apto)
        print()

        print('Valor Total Numérico do Imóvel:', i_valor_total_n)
        print()

        print('Valor Total Por Extenso do Imóvel:', i_valor_total_e)
        print()

        # 4.1 Texto

        texto = document.add_paragraph()
        texto.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        texto = texto.add_run(
            '\nRecebi de ' + i_pagador + ', CPF: ' + i_cpf + ' a importância de R$ ' + i_valor_numerico + ' (' + i_valor_extenso + ') referente ao sinal de entrada do APTO. ' + i_apto + ' do JPS 391 Smart Concept, adquirido no preço total de R$ ' + i_valor_total_n + ' (' + i_valor_total_e + '), a ser construído na cidade de Patos/PB.')
        fonte_do_texto = texto.font
        fonte_do_texto.name = 'Calibri'
        fonte_do_texto.size = Pt(14)

        # 5.0 Data - Entrada(s)

        print('Dia do Recibo:', i_dia)
        print()

        print('Mês do Recibo:', i_mes)
        print()

        print('Ano do Recibo:', i_ano)
        print()

        # 5.1 Data

        data = document.add_paragraph()
        data.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        data = data.add_run('\n' + i_dia + ' de ' + i_mes + ' de ' + i_ano + '\n')
        fonte_do_data = data.font
        fonte_do_data.name = 'Calibri'
        fonte_do_data.size = Pt(14)

        # 6.0 Assinatura

        assinatura = document.add_paragraph()
        assinatura.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        assinatura = assinatura.add_run(
            '\n____________________________________________\nJPS CONST. INCORP. E LOTEAMENTOS IMOBILIÁRIOS LTDA\nCNPJ: 38.484.970/0001-82')
        fonte_do_assinatura = assinatura.font
        fonte_do_assinatura.name = 'Calibri'
        fonte_do_assinatura.size = Pt(14)

        # 7.0 Rodapé

        section_f = document.sections[0]
        footer = section_f.footer

        # 7.1 Rodapé: Margens

        paragraph_f = footer.paragraphs[0]
        footer = section_f.footer
        section_f.footer_distance = Inches(0.1)
        paragraph_f.paragraph_format.left_indent = -Inches(1.14)

        # 7.2 Rodapé: Imagem

        run_f = paragraph_f.add_run()
        #run_f.add_picture("### ENDEREÇO DA IMAGEM ###")

        print('Nome do Arquivo:', i_impressao)

        document.save(i_impressao + '.docx')

class programa(funcoes):
    def __init__(self):
        self.root = root
        self.tela()
        self.frame_tela()
        self.apps_da_pagina_1()
        root.mainloop()

    def tela(self):
        self.root.title("Recibo de Pagamento")
        self.root.configure(background='#F0F0F0')
        self.root.geometry('400x400')
        self.root.resizable(True, True)
        self.root.maxsize(width=1920, height=1080)
        self.root.minsize(width=1, height=1)

    def frame_tela(self):
        self.frame = Frame(self.root, bd=0.1, bg='#FFFFFF',
                           highlightbackground='#F0F0F0', highlightthickness=2)
        self.frame.place(relx=0.005, rely=0.01, relwidth=0.99, relheight=0.98)

    def apps_da_pagina_1(self):

        # Configurações de Abas

        self.abas = ttk.Notebook(self.frame)

        self.aba_1 = Frame(self.abas)
        self.aba_1.configure(background='#F0F0F0')
        self.abas.add(self.aba_1, text=" Recibo ")

        self.abas.place(relx=0, rely=0, relwidth=1, relheight=1)

        self.aba_1_funcoes()

    def aba_1_funcoes(self):

        # 1.0 Fundo Geral

        self.fundo_geral = Label(self.aba_1, text='', relief="raised", bg='#F0F0F0', fg='#800000')
        self.fundo_geral.place(relx=0.005, rely=0.01, relwidth=0.985, relheight=0.98)

        # 2.0 Valor Numérico do Recibo

        self.fundo_01 = Label(self.aba_1, text='', relief="groove", bg='#F0F0F0', fg='#800000')
        self.fundo_01.place(relx=0.0145, rely=0.04, relwidth=0.48, relheight=0.13)

        self.valor_numerico_do_recibo_texto = Label(self.aba_1, text='Valor Numérico do Recibo :', bg='#F0F0F0',
                                               fg='#000000')
        self.valor_numerico_do_recibo_texto.place(relx=0.02, rely=0.03, relwidth=0.4, relheight=0.024)

        self.valor_numerico_do_recibo = Entry(self.aba_1, relief="sunken", bg='#FFFFFF', fg='#000000')
        self.valor_numerico_do_recibo.place(relx=0.025, rely=0.085, relwidth=0.4585, relheight=0.06)

        # 2.1 Valor Por Extenso do Recibo

        self.fundo_02 = Label(self.aba_1, text='', relief="groove", bg='#F0F0F0', fg='#800000')
        self.fundo_02.place(relx=0.0145, rely=0.2, relwidth=0.48, relheight=0.13)

        self.valor_por_extenso_do_recibo_texto = Label(self.aba_1, text='Valor Por Extenso do Recibo :', bg='#F0F0F0',
                                               fg='#000000')
        self.valor_por_extenso_do_recibo_texto.place(relx=0.02, rely=0.19, relwidth=0.425, relheight=0.026)

        self.valor_por_extenso_do_recibo = Entry(self.aba_1, relief="sunken", bg='#FFFFFF', fg='#000000')
        self.valor_por_extenso_do_recibo.place(relx=0.025, rely=0.245, relwidth=0.4585, relheight=0.06)

        # 2.2 Pagador

        self.fundo_03 = Label(self.aba_1, text='', relief="groove", bg='#F0F0F0', fg='#800000')
        self.fundo_03.place(relx=0.0145, rely=0.36, relwidth=0.48, relheight=0.13)

        self.pagador_texto = Label(self.aba_1, text='Pagador :', bg='#F0F0F0',
                                               fg='#000000')
        self.pagador_texto.place(relx=0.02, rely=0.34, relwidth=0.15, relheight=0.05)

        self.pagador = Entry(self.aba_1, relief="sunken", bg='#FFFFFF', fg='#000000')
        self.pagador.place(relx=0.025, rely=0.405, relwidth=0.4585, relheight=0.06)

        # 2.3 CPF

        self.fundo_04 = Label(self.aba_1, text='', relief="groove", bg='#F0F0F0', fg='#800000')
        self.fundo_04.place(relx=0.0145, rely=0.52, relwidth=0.48, relheight=0.13)

        self.cpf_texto = Label(self.aba_1, text='CPF :', bg='#F0F0F0',
                                               fg='#000000')
        self.cpf_texto.place(relx=0.02, rely=0.51, relwidth=0.1, relheight=0.026)

        self.cpf = Entry(self.aba_1, relief="sunken", bg='#FFFFFF', fg='#000000')
        self.cpf.place(relx=0.025, rely=0.565, relwidth=0.4585, relheight=0.06)

        # 2.4 Apartamento

        self.fundo_05 = Label(self.aba_1, text='', relief="groove", bg='#F0F0F0', fg='#800000')
        self.fundo_05.place(relx=0.0145, rely=0.68, relwidth=0.48, relheight=0.13)

        self.apartamento_texto = Label(self.aba_1, text='Apartamento :', bg='#F0F0F0',
                                               fg='#000000')
        self.apartamento_texto.place(relx=0.02, rely=0.66, relwidth=0.22, relheight=0.05)

        self.apartamento = Entry(self.aba_1, relief="sunken", bg='#FFFFFF', fg='#000000')
        self.apartamento.place(relx=0.025, rely=0.725, relwidth=0.4585, relheight=0.06)

        # 3.0 Valor Numérico Total

        self.fundo_06 = Label(self.aba_1, text='', relief="groove", bg='#F0F0F0', fg='#800000')
        self.fundo_06.place(relx=0.5, rely=0.04, relwidth=0.48, relheight=0.13)

        self.valor_numerico_total_texto = Label(self.aba_1, text='Valor Numérico Total :', bg='#F0F0F0',
                                                    fg='#000000')
        self.valor_numerico_total_texto.place(relx=0.5055, rely=0.03, relwidth=0.32, relheight=0.03)

        self.valor_numerico_total = Entry(self.aba_1, relief="sunken", bg='#FFFFFF', fg='#000000')
        self.valor_numerico_total.place(relx=0.5105, rely=0.085, relwidth=0.4585, relheight=0.06)

        # 3.1 Valor Por Extenso Total

        self.fundo_07 = Label(self.aba_1, text='', relief="groove", bg='#F0F0F0', fg='#800000')
        self.fundo_07.place(relx=0.5, rely=0.2, relwidth=0.48, relheight=0.13)

        self.valor_por_extenso_total_texto = Label(self.aba_1, text='Valor Por Extenso Total :', bg='#F0F0F0',
                                                    fg='#000000')
        self.valor_por_extenso_total_texto.place(relx=0.5055, rely=0.19, relwidth=0.35, relheight=0.026)

        self.valor_por_extenso_total = Entry(self.aba_1, relief="sunken", bg='#FFFFFF', fg='#000000')
        self.valor_por_extenso_total.place(relx=0.5105, rely=0.245, relwidth=0.4585, relheight=0.06)

        # 3.2 Dia

        self.fundo_08 = Label(self.aba_1, text='', relief="groove", bg='#F0F0F0', fg='#800000')
        self.fundo_08.place(relx=0.5, rely=0.36, relwidth=0.48, relheight=0.13)

        self.dia_texto = Label(self.aba_1, text='Dia :', bg='#F0F0F0',
                                                    fg='#000000')
        self.dia_texto.place(relx=0.5055, rely=0.34, relwidth=0.08, relheight=0.05)

        self.dia = Entry(self.aba_1, relief="sunken", bg='#FFFFFF', fg='#000000')
        self.dia.place(relx=0.5105, rely=0.405, relwidth=0.4585, relheight=0.06)

        # 3.3 Mês

        self.fundo_09 = Label(self.aba_1, text='', relief="groove", bg='#F0F0F0', fg='#800000')
        self.fundo_09.place(relx=0.5, rely=0.52, relwidth=0.48, relheight=0.13)

        self.mes_texto = Label(self.aba_1, text='Mês :', bg='#F0F0F0',
                                                    fg='#000000')
        self.mes_texto.place(relx=0.5055, rely=0.51, relwidth=0.1, relheight=0.026)

        self.mes = Entry(self.aba_1, relief="sunken", bg='#FFFFFF', fg='#000000')
        self.mes.place(relx=0.5105, rely=0.565, relwidth=0.4585, relheight=0.06)

        # 3.4 Ano

        self.fundo_10 = Label(self.aba_1, text='', relief="groove", bg='#F0F0F0', fg='#800000')
        self.fundo_10.place(relx=0.5, rely=0.68, relwidth=0.48, relheight=0.13)

        self.ano_texto = Label(self.aba_1, text='Ano :', bg='#F0F0F0',
                                                fg='#000000')
        self.ano_texto.place(relx=0.5055, rely=0.66, relwidth=0.1, relheight=0.05)

        self.ano = Entry(self.aba_1, relief="sunken", bg='#FFFFFF', fg='#000000')
        self.ano.place(relx=0.5105, rely=0.725, relwidth=0.4585, relheight=0.06)

        # 4.0 Número da Impressão

        self.fundo_11 = Label(self.aba_1, text='', relief="groove", bg='#F0F0F0', fg='#800000')
        self.fundo_11.place(relx=0.015, rely=0.855, relwidth=0.965, relheight=0.11)

        self.imp_texto = Label(self.aba_1, text='Nome do Arquivo :', bg='#F0F0F0',
                                               fg='#000000')
        self.imp_texto.place(relx=0.025, rely=0.888, relwidth=0.28, relheight=0.04)

        self.imp = Entry(self.aba_1, relief="sunken", bg='#FFFFFF', fg='#000000')
        self.imp.place(relx=0.315, rely=0.885, relwidth=0.299, relheight=0.05)

        # 5.0 Botões

        self.botao_gerar = tk.Button(self.aba_1, text='GERAR', bg='#F0F0F0', fg='#000000', command=self.gerar_recibo)
        self.botao_gerar.place(relx=0.637, rely=0.878, relwidth=0.15, relheight=0.06)

        self.botao_fechar = tk.Button(self.aba_1, text='FECHAR', bg='#F0F0F0', fg='#000000',command=self.sair)
        self.botao_fechar.place(relx=0.81, rely=0.878, relwidth=0.15, relheight=0.06)

programa()
